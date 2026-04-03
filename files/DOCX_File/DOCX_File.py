import html
import io
import urllib.error
import urllib.request


def _load_docx_bytes(path: str) -> bytes:
    if path.startswith("/mount/") or path.startswith("gdrive://"):
        import fsspec

        with fsspec.open(path, "rb") as f:
            return f.read()
    signed_url = fused.api.sign_url(path)
    req = urllib.request.Request(signed_url, method="GET")
    with urllib.request.urlopen(req, timeout=120) as resp:
        return resp.read()


def _docx_bytes_to_plain_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines)


@fused.udf(cache_max_age="30m")
def udf(path: str):
    try:
        raw = _load_docx_bytes(path)
        text = _docx_bytes_to_plain_text(raw)
    except urllib.error.URLError as e:
        safe = html.escape(str(e), quote=True)
        return f"""
    <!DOCTYPE html>
    <html>
    <head><title>DOCX Viewer</title></head>
    <body style="margin:0; padding:24px; font-family: system-ui, sans-serif; background:#1a1a1a; color:#ccc;">
        <h2 style="color:#ff6b6b;">Error loading DOCX</h2>
        <p>{safe}</p>
    </body>
    </html>
    """
    except Exception as e:
        safe = html.escape(str(e), quote=True)
        return f"""
    <!DOCTYPE html>
    <html>
    <head><title>DOCX Viewer</title></head>
    <body style="margin:0; padding:24px; font-family: system-ui, sans-serif; background:#1a1a1a; color:#ccc;">
        <h2 style="color:#ff6b6b;">Error parsing DOCX</h2>
        <p>{safe}</p>
    </body>
    </html>
    """

    body = html.escape(text)
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>DOCX Viewer</title>
        <style>
            * {{ margin: 0; padding: 0; box-sizing: border-box; }}
            body {{
                font-family: 'JetBrains Mono', 'Fira Code', 'Monaco', 'Consolas', monospace;
                background: #1a1a1a;
                color: #cccccc;
                height: 100vh;
                overflow: hidden;
            }}
            .text-content {{
                background: #1a1a1a;
                color: #cccccc;
                padding: 20px 20px 0 20px;
                height: 100vh;
                width: 100vw;
                overflow-y: auto;
                font-size: 15px;
                white-space: pre-wrap;
                word-wrap: break-word;
                line-height: 1.6;
            }}
            .text-content::-webkit-scrollbar {{ width: 8px; }}
            .text-content::-webkit-scrollbar-track {{ background: #1a1a1a; }}
            .text-content::-webkit-scrollbar-thumb {{ background: #D1E550; border-radius: 4px; }}
            .text-content::-webkit-scrollbar-thumb:hover {{ background: #E8FF59; }}
        </style>
    </head>
    <body>
        <div class="text-content">{body}</div>
    </body>
    </html>
    """
